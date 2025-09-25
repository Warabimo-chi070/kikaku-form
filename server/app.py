from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from datetime import datetime as dt, date
import os
from docx import Document

app = Flask(__name__)
CORS(app)

BASE = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = os.path.join(BASE, 'templates', '衣笠クラブ企画書フォーマット.docx')
OUT = os.path.abspath(os.path.join(BASE, '..', 'output'))
os.makedirs(OUT, exist_ok=True)

# ---------------------------
# ユーティリティ（置換）
# ---------------------------

def _replace_in_paragraph(paragraph, mapping):
    """段落内の複数Runに分割されたテキストも含め、まとめて置換して再構成"""
    if not mapping:
        return
    # 現在のrunsを一旦結合
    full_text = ''.join(run.text for run in paragraph.runs)
    replaced = full_text
    changed = False
    for k, v in mapping.items():
        if k in replaced:
            replaced = replaced.replace(k, v)
            changed = True
    if changed:
        # 既存のrunsをクリアして、置換後テキストの単一Runで再構成
        for _ in range(len(paragraph.runs)):
            paragraph.runs[0].clear()
            paragraph._p.remove(paragraph.runs[0]._r)
        paragraph.add_run(replaced)

def _replace_in_table(table, mapping):
    for row in table.rows:
        for cell in row.cells:
            _replace_in_block(cell, mapping)

def _replace_in_block(container, mapping):
    for p in container.paragraphs:
        _replace_in_paragraph(p, mapping)
    for t in container.tables:
        _replace_in_table(t, mapping)

def replace_placeholders(doc: Document, mapping: dict):
    # 文書直下
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    for t in doc.tables:
        _replace_in_table(t, mapping)
    # ヘッダ・フッタ側にプレースホルダがある場合（必要なら）
    for section in doc.sections:
        header = section.header
        if header:
            _replace_in_block(header, mapping)
        footer = section.footer
        if footer:
            _replace_in_block(footer, mapping)

# ---------------------------
# 日時の自動フォーマット
# ---------------------------

def build_formatted_datetime(y, m, d, ts, te, fallback=None):
    if not (y and m and d and ts and te):
        return fallback or ''
    try:
        yy, mm, dd = int(y), int(m), int(d)
        week = ['月','火','水','木','金','土','日']
        wd = week[date(yy, mm, dd).weekday()]  # Monday=0 -> 月
        ts2, te2 = str(ts)[:5], str(te)[:5]
        return f"{yy}年{mm}月{dd}日（{wd}） {ts2}-{te2}"
    except Exception:
        return fallback or ''

# ---------------------------
# メインAPI
# ---------------------------

@app.get('/template')
def get_template():
    """テンプレ全文をテキストとして返す（プレビュー用）"""
    if not os.path.exists(TEMPLATE):
        return jsonify({'ok': False, 'error': 'テンプレートがありません'}), 404

    doc = Document(TEMPLATE)
    lines = []

    # 段落をすべて収集
    for p in doc.paragraphs:
        lines.append(p.text)

    # 表も収集（セルをタブ区切り）
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text for c in row.cells]
            lines.append("\t".join(cells))

    return jsonify({
        'ok': True,
        'template_text': "\n".join(lines)
    })

@app.post('/generate')
def generate():
    data = request.get_json(force=True) or {}
    now = dt.now().strftime('%Y%m%d_%H%M%S')
    title = data.get('title') or '無題企画'
    out_name = f'企画書_{title}_{now}.docx'
    out_path = os.path.join(OUT, out_name)

    # 1) テンプレ読み込み（存在チェック）
    if not os.path.exists(TEMPLATE):
        return jsonify({
            'ok': False,
            'error': f'テンプレートが見つかりません: {TEMPLATE}',
            'hint': 'server/templates/ に衣笠クラブ企画書フォーマット.docx を配置してください。'
        }), 400
    doc = Document(TEMPLATE)

    # 2) 日時（自動整形）
    formatted_dt = data.get('formatted_datetime') or build_formatted_datetime(
        data.get('year'), data.get('month'), data.get('day'),
        data.get('time_start'), data.get('time_end'),
        fallback=data.get('datetime')
    )

    # 3) マッピング（{{token}} → 値）
    skills = data.get('skills') or []
    mapping = {
        '{{title}}': data.get('title', '') or '',
        '{{club}}': data.get('club', '') or '',
        '{{dept}}': data.get('dept', '') or '',
        '{{category}}': data.get('category', '') or '',
        '{{field}}': data.get('field', '') or '',
        '{{datetime}}': formatted_dt or '',
        '{{place}}': data.get('place', '') or '',
        '{{expected_ivusa}}': str(data.get('expected_ivusa') or ''),
        '{{expected_other}}': str(data.get('expected_other') or ''),
        '{{owner}}': data.get('owner', '') or '',
        '{{other_club}}': data.get('other_club', '') or '',
        '{{beneficiary}}': data.get('beneficiary', '') or '',
        '{{is_new}}': data.get('is_new', '') or '',
        '{{activity_kind}}': data.get('activity_kind', '') or '',
        '{{duration}}': data.get('duration', '') or '',
        '{{skills}}': '、'.join(skills) if skills else '',
        '{{purpose}}': data.get('purpose', '') or '',
        '{{kpi}}': data.get('kpi', '') or '',
        '{{details}}': data.get('details', '') or '',
        '{{pre_schedule}}': data.get('pre_schedule', '') or '',
        '{{exec_schedule}}': data.get('exec_schedule', '') or '',
        '{{day_schedule}}': data.get('day_schedule', '') or '',
        '{{risk_before}}': data.get('risk_before', '') or '',
        '{{risk_during}}': data.get('risk_during', '') or '',
        '{{stakeholders}}': data.get('stakeholders', '') or '',
        '{{permit_city}}': data.get('permit_city', '') or '',
        '{{permit_fire}}': data.get('permit_fire', '') or '',
        '{{permit_police}}': data.get('permit_police', '') or '',
        '{{other_notes}}': data.get('other_notes', '') or '',
        '{{budget_total}}': str(data.get('budget_total') or ''),
        '{{funding}}': data.get('funding', '') or '',
        '{{budget_usage}}': data.get('budget_usage', '') or '',
        '{{press}}': data.get('press', '') or '',
        '{{drive_student}}': data.get('drive_student', '') or '',
        '{{rentacar}}': data.get('rentacar', '') or '',
        '{{general_join}}': data.get('general_join', '') or '',
        '{{knife}}': data.get('knife', '') or '',
        '{{power_tool}}': data.get('power_tool', '') or '',
        '{{self_cook}}': data.get('self_cook', '') or '',
        '{{stay}}': data.get('stay', '') or '',
        '{{emg_clubman}}': data.get('emg_clubman', '') or '',
        '{{emg_officer}}': data.get('emg_officer', '') or '',
        '{{emg_day}}': data.get('emg_day', '') or '',
        '{{emg_ok}}': data.get('emg_ok', '') or '',
    }

    # 4) 置換（段落・表・ヘッダ・フッタを横断）
    replace_placeholders(doc, mapping)

    # 5) 保存
    doc.save(out_path)

    return jsonify({
        'ok': True,
        'output_path': out_path,
        'download_url': f'/download/{os.path.basename(out_path)}'
    })

@app.get('/download/<path:fname>')
def download(fname):
    return send_from_directory(OUT, fname, as_attachment=True)

@app.get('/health')
def health():
    return jsonify({'ok': True})

if __name__ == '__main__':
    # 127.0.0.1固定：フロントのserverURLに合わせています
    app.run(host='127.0.0.1', port=5000, debug=False)